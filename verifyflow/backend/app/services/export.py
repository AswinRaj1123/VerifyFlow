from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from sqlalchemy.orm import Session
from datetime import datetime
import io
from typing import BinaryIO


def build_docx(session_id: int, session_title: str, questions: list, db: Session) -> BinaryIO:
    """
    Build a Word document with questionnaire structure:
    - Session title as heading
    - Each question with answer and citations
    - Edited flag if applicable
    
    Returns a BytesIO object ready for download.
    """
    doc = Document()
    
    # Add session title as Heading 1
    title = doc.add_heading(session_title or f"Questionnaire Session {session_id}", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add metadata
    meta_para = doc.add_paragraph()
    meta_para.add_run(f"Generated: {datetime.utcnow().strftime('%B %d, %Y at %H:%M UTC')}").italic = True
    meta_para.add_run(f"\nTotal Questions: {len(questions)}").italic = True
    meta_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add separator
    doc.add_paragraph("_" * 80)
    doc.add_paragraph()
    
    # Process each question
    for idx, question in enumerate(questions, 1):
        # Question number and text (bold)
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(f"Q{idx}. {question.original_text}")
        q_run.bold = True
        q_run.font.size = Pt(12)
        q_run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        
        # Answer text (normal)
        answer_text = question.answer if question.answer else "Not answered yet."
        answer_para = doc.add_paragraph(answer_text)
        answer_para.paragraph_format.left_indent = Inches(0.5)
        answer_para.paragraph_format.space_after = Pt(6)
        
        # Confidence score
        if question.confidence > 0:
            conf_para = doc.add_paragraph()
            conf_para.paragraph_format.left_indent = Inches(0.5)
            conf_run = conf_para.add_run(f"Confidence: {question.confidence}%")
            conf_run.font.size = Pt(9)
            
            # Color code confidence
            if question.confidence >= 80:
                conf_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
            elif question.confidence >= 50:
                conf_run.font.color.rgb = RGBColor(255, 140, 0)  # Orange
            else:
                conf_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
        
        # Citations (italic, small, indented)
        if question.citations and len(question.citations) > 0:
            cit_para = doc.add_paragraph()
            cit_para.paragraph_format.left_indent = Inches(0.5)
            cit_run = cit_para.add_run("Citations:")
            cit_run.italic = True
            cit_run.font.size = Pt(9)
            cit_run.font.color.rgb = RGBColor(100, 100, 100)  # Gray
            
            for citation in question.citations:
                cit_item = doc.add_paragraph(f"• {citation}", style='List Bullet')
                cit_item.paragraph_format.left_indent = Inches(1.0)
                cit_item.runs[0].font.size = Pt(9)
                cit_item.runs[0].italic = True
                cit_item.runs[0].font.color.rgb = RGBColor(100, 100, 100)
        
        # Edited flag
        if question.is_edited:
            edit_para = doc.add_paragraph()
            edit_para.paragraph_format.left_indent = Inches(0.5)
            edit_run = edit_para.add_run("✏️ Edited by user")
            edit_run.font.size = Pt(8)
            edit_run.font.color.rgb = RGBColor(150, 150, 150)  # Light gray
            edit_run.italic = True
        
        # Add spacing between questions
        doc.add_paragraph()
    
    # Add footer with session info
    doc.add_paragraph("_" * 80)
    footer_para = doc.add_paragraph()
    footer_run = footer_para.add_run(f"Session ID: {session_id} | Exported from VerifyFlow")
    footer_run.font.size = Pt(8)
    footer_run.italic = True
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Save to BytesIO
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return file_stream
