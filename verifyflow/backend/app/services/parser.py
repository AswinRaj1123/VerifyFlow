import pdfplumber
import pandas as pd
from docx import Document
from openai import OpenAI
import json
from io import BytesIO

client = OpenAI()

def parse_questionnaire(file_bytes: bytes, filename: str) -> list[str]:
    text = ""
    if filename.endswith(".pdf"):
        with pdfplumber.open(BytesIO(file_bytes)) as pdf:
            text = "\n".join([p.extract_text() or "" for p in pdf.pages])
    elif filename.endswith((".xlsx", ".xls")):
        df = pd.read_excel(BytesIO(file_bytes))
        return df.iloc[:, 0].dropna().tolist()  # assume first column = questions
    elif filename.endswith(".docx"):
        doc = Document(BytesIO(file_bytes))
        text = "\n".join([p.text for p in doc.paragraphs])

    # AI fallback for messy PDFs
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "system", "content": "Extract every question as a clean JSON array of strings."},
                  {"role": "user", "content": f"Text:\n{text}"}],
        response_format={"type": "json_object"}
    )
    return json.loads(response.choices[0].message.content)["questions"]


def parse_reference(file_bytes: bytes, filename: str) -> str:
    """
    Extract clean, readable text from PDF, DOCX, TXT.
    Returns concatenated text (we chunk later in RAG).
    """
    text = ""
    try:
        if filename.lower().endswith(".pdf"):
            with pdfplumber.open(BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
        elif filename.lower().endswith(".docx"):
            doc = Document(BytesIO(file_bytes))
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            # optional: also handle tables if needed later
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text.strip() + " | "
                    text += "\n"
        elif filename.lower().endswith(".txt"):
            text = file_bytes.decode("utf-8", errors="replace")
        else:
            raise ValueError("Unsupported file type")

        # Minimal cleanup — remove excessive newlines
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        return "\n".join(lines)
    except Exception as e:
        raise RuntimeError(f"Parsing failed: {str(e)}")